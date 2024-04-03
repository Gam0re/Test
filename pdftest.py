from docx import Document
import os
import comtypes.client

# Открываем исходный Word документ
docx_file = "first.docx"
doc = Document(docx_file)

# Заменяем кодовые слова в каждом абзаце документа
for para in doc.paragraphs:
    if "name" in para.text:
        para.text = para.text.replace("name", "Иванов Олег Васильевич")

# Заменяем кодовые слова в каждой ячейке таблицы документа
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            if "name" in cell.text:
                cell.text = cell.text.replace("name", "Иванов Олег Васильевич")

# Сохраняем измененный документ
doc.save("output.docx")

word_path = "output.docx"
pdf_path = "output.pdf"

doc = Document(word_path)

word = comtypes.client.CreateObject("Word.Application")
docx_path = os.path.abspath(word_path)
pdf_path = os.path.abspath(pdf_path)

pdf_format = 17
word.Visible = False
in_file = word.Documents.Open(docx_path)
in_file.SaveAs(pdf_path, FileFormat=pdf_format)
in_file.Close()

word.Quit()
